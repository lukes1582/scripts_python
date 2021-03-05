'''
Created on 05/03/2021
@author: lukes158@gmail.com l0m1s
'''
from docx import Document
from docx.shared import Cm


document = Document()

document.add_heading('Prova primo documento docx', 0)

p = document.add_paragraph('questo è un paragrafo')
p.add_run('bold').bold = True
p.add_run(' e anche  ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph('primo elemento di una lista a punti', style='List Bullet')
document.add_paragraph('primo elemento di una lista numerata', style='List Number')

document.add_picture('img.png', width=Cm(10.25))

records = (
    (3, '105', 'prezzemolo'),
    (7, '422', 'aglio'),
    (4, '638', 'Olio, pepe, peperoncino, e cannella')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Quantità'
hdr_cells[1].text = 'Identificativo'
hdr_cells[2].text = 'Descrizione'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('primo-demo.docx')